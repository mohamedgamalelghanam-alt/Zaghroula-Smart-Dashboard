import streamlit as st
import pandas as pd
import plotly.express as px
from docx import Document
from io import BytesIO
from datetime import datetime
import re

# إعدادات الصفحة - Dark Mode مريح وواضح
st.set_page_config(page_title="Zaghloula Dashboard", layout="wide")

st.markdown("""
<style>
    .stApp { background-color: #0e1117; color: #ffffff; }
    .metric-card {
        background-color: #1a1c23; padding: 20px; border-radius: 12px;
        border: 1px solid #30363d; text-align: center; margin-bottom: 10px;
    }
    .sales-val { color: #2ecc71; font-size: 2.2rem; font-weight: bold; }
    .profit-val { color: #3498db; font-size: 2.2rem; font-weight: bold; }
    .warning-val { color: #e74c3c; font-size: 2.2rem; font-weight: bold; }
</style>
""", unsafe_allow_html=True)

def create_word_report(data, t_sales, t_profit, branch):
    doc = Document()
    doc.add_heading(f'تقرير زغلولة - {branch}', 0)
    table = doc.add_table(rows=1, cols=3); table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = 'الصنف', 'المبيعات', 'الربح'
    for _, row in data.head(20).iterrows():
        c = table.add_row().cells
        c[0].text, c[1].text, c[2].text = str(row['الصنف']), f"{row.get('Total_Sales', 0):,.2f}", f"{row.get('Net_Profit', 0):,.2f}"
    target = BytesIO(); doc.save(target)
    return target.getvalue()

st.title("📊 نظام زغلولة لإدارة البيانات")

uploaded_file = st.sidebar.file_uploader("📂 ارفع ملف (Excel Workbook)", type=['xlsx', 'xls', 'csv'])

if uploaded_file:
    try:
        # قراءة الملف بذكاء
        if uploaded_file.name.endswith('csv'):
            try: df = pd.read_csv(uploaded_file, encoding='cp1256')
            except: df = pd.read_csv(uploaded_file, encoding='utf-8-sig')
        else:
            df = pd.read_excel(uploaded_file)

        # تنظيف الداتا: مسح الصفوف الفاضية تماماً
        df = df.dropna(how='all').reset_index(drop=True)
        
        if df.empty:
            st.error("⚠️ الملف ده فاضي يا هندسة، تأكد إنك حفظت البيانات فيه.")
        else:
            # محاولة إيجاد الأعمدة بذكاء لو الأسامي متغيرة
            cols = df.columns.tolist()
            col_item = next((c for c in cols if 'صنف' in str(c)), cols[0])
            col_qty = next((c for c in cols if 'كمية' in str(c)), cols[min(1, len(cols)-1)])
            col_price = next((c for c in cols if 'سعر' in str(c)), cols[min(2, len(cols)-1)])
            col_stock = next((c for c in cols if 'متبقي' in str(c) or 'رصيد' in str(c)), None)

            # تنظيف الأرقام
            def to_num(x):
                try: return float(re.sub(r'[^\d.]', '', str(x)))
                except: return 0.0

            df['الصنف'] = df[col_item].astype(str)
            df['Qty'] = df[col_qty].apply(to_num)
            df['Price'] = df[col_price].apply(to_num)
            df['Stock'] = df[col_stock].apply(to_num) if col_stock else 10.0
            
            df['Total_Sales'] = df['Qty'] * df['Price']
            df['Net_Profit'] = df['Total_Sales'] * 0.25 # نسبة ربح تقديرية

            df_final = df[df['Total_Sales'] > 0].copy()
            branch = "فرع إيتاي"

            # العرض
            ts, tp = df_final['Total_Sales'].sum(), df_final['Net_Profit'].sum()
            low_stock = df_final[df_final['Stock'] <= 5]

            c1, c2, c3 = st.columns(3)
            with c1: st.markdown(f'<div class="metric-card"><h2>💰 مبيعات اليوم</h2><div class="sales-val">{ts:,.2f}</div></div>', unsafe_allow_html=True)
            with c2: st.markdown(f'<div class="metric-card"><h2>📈 الأرباح التقديرية</h2><div class="profit-val">{tp:,.2f}</div></div>', unsafe_allow_html=True)
            with c3: st.markdown(f'<div class="metric-card"><h2>🚨 نواقص المخزن</h2><div class="warning-val">{len(low_stock)}</div></div>', unsafe_allow_html=True)

            t1, t2 = st.tabs(["🛒 قائمة النواقص", "📊 التحليل"])
            with t1:
                st.subheader("📦 بضاعة لازم تطلبها")
                st.dataframe(low_stock[['الصنف', 'Stock']].rename(columns={'Stock': 'الرصيد'}), use_container_width=True)
            with t2:
                st.plotly_chart(px.bar(df_final.nlargest(10, 'Total_Sales'), x='Total_Sales', y='الصنف', orientation='h', template="plotly_dark"))

    except Exception as e:
        st.error(f"⚠️ حصلت مشكلة: {e}. جرب تحفظ الملف بصيغة Excel Workbook وارفعُه تاني.")
else:
    st.info("مستني ترفع ملف المبيعات يا هندسة..")

# دالة إنشاء تقرير الوورد
def create_word_report(data, t_sales, t_profit, branch):
    doc = Document()
    doc.add_heading(f'تقرير مبيعات: {branch}', 0)
    doc.add_paragraph(f'التاريخ: {datetime.now().strftime("%Y-%m-%d")}')
    table = doc.add_table(rows=1, cols=3); table.style = 'Table Grid'
    hdr = table.rows[0].cells; hdr[0].text, hdr[1].text, hdr[2].text = 'الصنف', 'المبيعات', 'الربح'
    for _, row in data.head(20).iterrows():
        c = table.add_row().cells
        c[0].text, c[1].text, c[2].text = str(row['الصنف']), f"{row['Total_Sales']:,.2f}", f"{row['Net_Profit']:,.2f}"
    target = BytesIO(); doc.save(target)
    return target.getvalue()

st.title("📊تحليل مبيعات زغلوله اليومي")
st.markdown("---")

uploaded_file = st.sidebar.file_uploader("📂 ارفع ملف المبيعات", type=['xls', 'csv', 'xlsx'])

if uploaded_file:
    try:
        # قراءة الملف بأمان
        try: df = pd.read_excel(uploaded_file)
        except:
            uploaded_file.seek(0)
            try: df = pd.read_csv(uploaded_file, encoding='cp1256')
            except: 
                uploaded_file.seek(0)
                df = pd.read_csv(uploaded_file, encoding='utf-8-sig')

        df.columns = [str(c).strip() for c in df.columns]
        branch = df['الفرع'].dropna().iloc[0] if 'الفرع' in df.columns else "زغلولة - فرع إيتاي"

        # تنظيف الأرقام من الرموز (×, *)
        def clean_num(x):
            val = re.sub(r'[^\d.]', '', str(x))
            try: return float(val) if val else 0
            except: return 0

        # ربط الأعمدة
        col_item = 'الصنف' if 'الصنف' in df.columns else df.columns[2]
        col_qty = 'الكمية' if 'الكمية' in df.columns else df.columns[3]
        col_price = 'السعر' if 'السعر' in df.columns else df.columns[4]
        col_stock = 'الكمية المتبقية' if 'الكمية المتبقية' in df.columns else None

        df['الصنف'] = df[col_item].astype(str)
        df['Qty'] = df[col_qty].apply(clean_num)
        df['Price'] = df[col_price].apply(clean_num)
        df['Stock'] = df[col_stock].apply(clean_num) if col_stock else 10
        
        df['Total_Sales'] = df['Qty'] * df['Price']
        df['Net_Profit'] = df['Total_Sales'] * 0.25 # افتراضي 25%
        
        df_final = df[df['Total_Sales'] > 0].copy()
        df_final = df_final[~df_final['الصنف'].str.contains('اجمالى|وارد', na=False)]

        # حساب المؤشرات
        ts, tp = df_final['Total_Sales'].sum(), df_final['Net_Profit'].sum()
        low_stock = df_final[df_final['Stock'] <= 5]

        # --- عرض الكروت بوضوح عالي جداً ---
        st.subheader(f"📍 مراجعة الأداء اليومي: {branch}")
        c1, c2, c3 = st.columns(3)
        with c1: st.markdown(f'<div class="metric-card"><h2>💰 إجمالي المبيعات</h2><h1 class="sales-val">{ts:,.2f}</h1></div>', unsafe_allow_html=True)
        with c2: st.markdown(f'<div class="metric-card"><h2>📈 صافي الأرباح</h2><h1 class="profit-val">{tp:,.2f}</h1></div>', unsafe_allow_html=True)
        with c3: st.markdown(f'<div class="metric-card"><h2>🚨 أصناف للنواقص</h2><h1 class="warning-val">{len(low_stock)}</h1></div>', unsafe_allow_html=True)

        st.markdown("<br>", unsafe_allow_html=True)

        # التبويبات
        tab1, tab2, tab3 = st.tabs(["📈 التحليل البياني", "📦 قائمة النواقص", "📄 التقارير"])

        with tab1:
            col_l, col_r = st.columns(2)
            with col_l:
                top_p = df_final.groupby('الصنف')['Total_Sales'].sum().nlargest(10).reset_index()
                st.plotly_chart(px.bar(top_p, x='Total_Sales', y='الصنف', orientation='h', title="أعلى الأصناف مبيعاً", template="plotly_dark"), use_container_width=True)
            with col_r:
                fig_pie = px.pie(df_final.head(10), values='Total_Sales', names='الصنف', title="توزيع مبيعات الأصناف", hole=0.4, template="plotly_dark")
                st.plotly_chart(fig_pie, use_container_width=True)

        with tab2:
            st.subheader("🛒 قائمة بضاعة النواقص (رصيد أقل من 5)")
            if not low_stock.empty:
                # عرض الجدول بخلفية واضحة وكلام أبيض صريح
                st.dataframe(low_stock[['الصنف', 'Stock']].drop_duplicates().rename(columns={'Stock': 'الكمية المتبقية'}), use_container_width=True)
            else:
                st.success("المخزن مكتمل، مفيش أي نواقص!")

        with tab3:
            st.markdown("### استخراج تقرير رسمي")
            st.info("اضغط على الزر أدناه لتحميل تقرير المبيعات بصيغة Word.")
            st.download_button("📥 تحميل التقرير (Word)", data=create_word_report(df_final, ts, tp, branch), file_name=f"تقرير_زغلولة_{datetime.now().strftime('%Y%m%d')}.docx")

    except Exception as e:
        st.error(f"حدث خطأ: {e}")
else:
    st.info("يا هندسة ارفع ملف المبيعات من القائمة الجانبية عشان نعرض لك الأرباح!")

st.markdown("---")
st.markdown("<center>تطوير المهندس محمد جمال | 2026</center>", unsafe_allow_html=True)
